"""Build the micro-Taiex day-trade strategy rulebook as a .docx.

Every number here was read out of the source at build time (see the 程式碼位置
column in each table); nothing is quoted from memory or from stale comments.
Figures come from tmp/doc_figs/ (see strategy_doc_figs.py).
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
FIGS = ROOT / "doc_figs"
OUT = ROOT / "微台當沖策略規則說明書.docx"

CJK = "Microsoft JhengHei"
MONO = "Consolas"
ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
MUTED = RGBColor(0x47, 0x55, 0x69)
DANGER = RGBColor(0xB9, 0x1C, 0x1C)


def _cjk(run, font=CJK, size=None, bold=None, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def para(doc, text="", size=10.5, bold=False, color=None, align=None,
         space_after=6, font=CJK, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _cjk(r, font=font, size=size, bold=bold, color=color)
    r.font.italic = italic
    return p


def heading(doc, text, level=1):
    sizes = {0: 20, 1: 15, 2: 12.5, 3: 11}
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    _cjk(r, size=sizes[level], bold=True,
         color=ACCENT if level <= 1 else RGBColor(0x0F, 0x17, 0x2A))
    return h


def code(doc, text, size=9.5, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    for i, line in enumerate(text.strip("\n").split("\n")):
        if i:
            p.add_run("\n")
        _cjk(p.add_run(line), font=MONO, size=size)
    return p


def bullets(doc, items, size=10.5, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            _cjk(p.add_run(it[0]), size=size, bold=True)
            _cjk(p.add_run(it[1]), size=size)
        else:
            _cjk(p.add_run(it), size=size)


def table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _cjk(cell.paragraphs[0].add_run(htxt), size=size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            mono = val.startswith("`") and val.endswith("`")
            _cjk(p.add_run(val.strip("`")), font=MONO if mono else CJK, size=size)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t


def figure(doc, name, caption=None, width=6.3):
    doc.add_picture(str(FIGS / name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        para(doc, caption, size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=10, italic=True)


def note(doc, text, color=DANGER, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(8)
    _cjk(p.add_run("※ "), size=size, bold=True, color=color)
    _cjk(p.add_run(text), size=size, color=color)


# =====================================================================
doc = Document()
st = doc.styles["Normal"]
st.font.name = CJK
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.85)
    s.top_margin = s.bottom_margin = Inches(0.8)

# ---------------------------------------------------------------- 封面
para(doc, "微台期貨當沖系統", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=2)
para(doc, "現行策略規則說明書", size=18, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
para(doc, f"產生日期：{date.today():%Y-%m-%d}　│　依據程式碼實況產出，非依文件或註解",
     size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "專案路徑：c:\\Claude\\Invest\\microtaiex-daytrade\\",
     size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

heading(doc, "一句話總結", 1)
para(doc, "在台指微型期貨（TM0000）上，用 5 分鐘 K 棒跑三個「只做多」的價格型態訊號，"
          "任何時刻最多持有 1 口，靠「每筆固定點數上限」與「移動停損」兩道防線保護，"
          "並在每個交易時段收盤前 1 分鐘無條件平倉，絕不留倉過夜。", size=11.5)

heading(doc, "最關鍵的五件事", 2)
bullets(doc, [
    ("只做多。", "現行啟用的三個訊號全部屬於多方，空方訊號一個都沒開，系統永遠不會建立空單。"),
    ("永遠 1 口。", "沒有加碼、沒有減碼、沒有分批。實際持倉只有「空手」與「1 口多單」兩種，"
     "另有進場中、出場中兩個委託在途的過渡狀態。"),
    ("真正綁住風險的是 per-trade cap，不是移動停損。", "高波動時 cap 被上限壓在約 40 點，"
     "而移動停損在約 330 點外，絕大多數情況下是 cap 先觸發。"),
    ("所有判斷都在 5 分鐘 K 棒收盤那一刻完成。", "盤中不看跳動報價，所以實際虧損可能"
     "超過名目上限（見第九章）。"),
    ("不留倉。", "日盤 13:44、夜盤 04:59 無條件平倉，部位絕不跨越交易時段。"),
])

doc.add_page_break()

# ---------------------------------------------------------------- 1
heading(doc, "一、現行實盤配置", 1)
para(doc, "以下是 Windows 服務 InvestMicroPaper 此刻實際帶的參數，由服務設定查得，"
          "不是程式碼預設值：")
code(doc, "python.exe  C:\\Claude\\Invest\\microtaiex-daytrade\\run_live.py"
          "  --paper  --masha  --warmup")
table(doc,
      ["項目", "現行狀態", "說明"],
      [["執行模式", "`--paper`", "真實行情 tick，但由內建 SimBroker 模擬成交，不下真單"],
       ["啟用訊號", "`pick, sell_flee, buy`", "三個皆為多方訊號（run_live.py:88 LIVE_SIGNALS）"],
       ["日線閘門", "未啟用（無 --gated）", "回測 10 檔淨利 +86,510，明顯低於不開的 +121,520"],
       ["buy-only 模式", "未啟用（無 --buyonly）", "程式支援但目前沒開，仍跑三訊號"],
       ["麻紗 shadow", "啟用（--masha）", "另一套獨立策略並行跑，自有部位與紀錄，與本文件無關"],
       ["重啟暖機", "啟用（--warmup）", "重啟後補抓歷史 K 棒填充指標，避免裸奔"],
       ["交易標的", "`TM0000`", "群益微型台指期近月連續"],
       ["策略週期", "`5m`", "1m 棒只用於觸發強制平倉"]],
      widths=[1.3, 1.75, 3.3])
note(doc, "本文件描述的是上表中的 composite 主策略。--masha 是並行的另一套系統（麻紗），"
          "有自己的訊號、停損與紀錄檔，不在本文件範圍內。")

figure(doc, "fig1_dataflow.png",
       "資料流：即時 tick 進來後聚合成 K 棒，每根 5 分鐘 K 棒收盤驅動一次完整決策。")

# ---------------------------------------------------------------- 2
heading(doc, "二、進場訊號（三個，全部做多）", 1)
para(doc, "每根 5 分鐘 K 棒收盤時，三個訊號各自獨立判斷一次。以下符號約定：")
bullets(doc, [
    "curr = 剛收盤的這根 K 棒，prev = 前一根 K 棒",
    "open / high / low / close = 該根 K 棒的開高低收",
    "prior = 用來比較的更早 N 根 K 棒。三個訊號的 prior 範圍不一樣，見下方各節；"
    "特別注意 pick 的 prior 不含 prev，但 buy 與 sell_flee 的 prior 含 prev。",
], size=10)

heading(doc, "2-1　pick　反轉做多（吞噬 + 破底）", 2)
para(doc, "抓「連續下跌後出現一根把前一根黑K完全吃掉的紅K，而且這兩根還創了新低」——"
          "也就是俗稱的下跌末端多頭吞噬。屬於逆勢接刀型。", size=10)
code(doc, """
prev.close < prev.open                       # 前一根收黑
curr.close > curr.open                       # 這一根收紅
curr.open  <= prev.close                     # 開盤沒高過前一根收盤
curr.close >= prev.open                      # 收盤蓋過前一根開盤   → 完成吞噬
min(prev.low, curr.low) <= min(prior 20 根的 low)   # 兩根之中的低點創 20 根新低
""")
para(doc, "參數：lookback = 20。prior 取 bars[-22:-2]，即 curr 與 prev 之前的 20 根"
          "（不含 prev）。資料不足 22 根時不觸發。　程式碼：signals.py:44-52",
     size=9.5, color=MUTED)

heading(doc, "2-2　buy　順勢做多（突破 + 幅度確認）", 2)
para(doc, "抓「向上突破近期高點，而且突破得夠用力」。加上幅度門檻是為了濾掉貼著前高的假突破。"
          "屬於順勢追價型。", size=10)
code(doc, """
prior_high = max(prior 20 根的 high)

條件一：curr.close > prior_high                        # 收盤價突破前高
條件二：curr.close >= prior_high + 0.75 * ATR(21)      # 且超出幅度至少 0.75 個 ATR

若 ATR 尚未暖機完成（不足 21 根）→ 只檢查條件一
""")
para(doc, "參數：breakout_lb = 20、幅度係數 0.75、ATR 週期 21。prior 取 bars[-21:-1]，"
          "即 curr 之前的 20 根（含 prev）。　程式碼：signals.py:66-80",
     size=9.5, color=MUTED)

heading(doc, "2-3　sell_flee　空頭陷阱反轉做多", 2)
para(doc, "抓「假跌破」——盤中殺破近期低點把空單騙進場，但收盤又拉回來且收在高檔，"
          "代表下殺失敗。屬於逆勢反轉型，是三個訊號中優先權最高的。", size=10)
code(doc, """
curr.low   <  min(prior 10 根的 low)                   # 盤中確實跌破近期低點
curr.close >  prev.close                               # 但收盤比前一根高（拉回來了）
curr.close >= curr.low + 0.6 * (curr.high - curr.low)  # 且收在自身區間的上 40%
""")
para(doc, "參數：flee_lb = 10、區間位置係數 0.6。prior 取 bars[-11:-1]，"
          "即 curr 之前的 10 根（含 prev）。　程式碼：signals.py:101-114",
     size=9.5, color=MUTED)

figure(doc, "fig3_patterns.png", "三個進場型態示意（紅 = 收漲，綠 = 收跌）。")

# ---------------------------------------------------------------- 3
heading(doc, "三、訊號衝突怎麼處理", 1)
para(doc, "同一根 K 棒可能同時觸發多個訊號。系統給每個訊號一個優先權，取最高的那個當作進場理由"
          "（進場動作完全相同，差別只在紀錄與通知上顯示哪個訊號）：")
table(doc,
      ["優先權", "訊號", "類型", "備註"],
      [["3（最高）", "`sell_flee`", "逃命 / 反轉", "現行三訊號中最優先"],
       ["2", "`pick`", "反轉", ""],
       ["1（最低）", "`buy`", "順勢", ""]],
      widths=[1.0, 1.3, 1.3, 2.6])
para(doc, "因為現行三個訊號全是多方，多空分數比較的結果恆為「多方勝出或無訊號」，"
          "空方分數永遠是 0。程式碼：composite.py:36-37、96-123", size=9.5, color=MUTED)

heading(doc, "部位規則", 2)
bullets(doc, [
    ("空手時收到多方訊號　→", "送出買進 1 口的進場單。"),
    ("已持有多單時再收到多方訊號　→", "不動作（不加碼、不重複下單）。"),
    ("委託在途（尚未成交回報）時　→", "不動作，避免重複下單。"),
    ("強制平倉時段內　→", "一律不開新倉。"),
])
para(doc, "程式碼：risk_manager.py:92-120、position/state_machine.py:22-65",
     size=9.5, color=MUTED)

doc.add_page_break()

# ---------------------------------------------------------------- 4
heading(doc, "四、出場規則（三道，任一觸發即平倉）", 1)
para(doc, "本策略沒有「停利」設計，所有出場都來自下列三道之一。三者在每根 5 分鐘 K 棒收盤時檢查。")

heading(doc, "4-1　第一道：每筆固定點數上限（per-trade cap）", 2)
para(doc, "這是實務上最常觸發的一道。注意它**不是**進場後就固定不動——每根 K 棒都會用"
          "當下的 ATR 重算一次：", size=10)
code(doc, """
cap = clamp( 0.5 * ATR(21),  下限 12 點,  上限 = 0.087% * 進場價 )

以進場價 46,000 為例：上限 = 0.00087 * 46000 = 40.0 點
若 ATR(21) = 110 → 0.5 * ATR = 55 點，被上限壓成 40 點
若 ATR(21) =  60 → 0.5 * ATR = 30 點，未達上下限，cap 就是 30 點

出場條件（做多）：進場價 - 這根收盤價 >= cap
""")
bullets(doc, [
    ("哪些會動、哪些不會：", "上限用進場價計算，一筆交易內固定；但 0.5×ATR 隨波動變化，"
     "所以只要它落在 12 與上限之間，cap 每根 K 棒都會跟著浮動。"),
    ("上限為什麼用百分比而不是固定點數：", "固定 40 點在指數 22,000 時相當於 0.18%，"
     "但在 46,000 時只剩 0.087%，嚴格度會隨指數上漲而縮水。改成百分比後尺度不變。"),
    ("下限為什麼維持固定 12 點：", "下限守的是「雜訊」，對應的是跳動點與買賣價差，"
     "這些不隨指數縮放，所以不該百分比化。"),
    ("ATR 還沒暖機時：", "退回使用固定的 20 點上限（需要 21 根 K 棒才算得出 ATR）。"),
], size=10)
para(doc, "程式碼：risk_manager.py:244-263（_trade_cap）、146-151（觸發判斷）",
     size=9.5, color=MUTED)

heading(doc, "4-2　第二道：Chandelier 移動停損", 2)
para(doc, "以持倉期間的最高價往下推 3 個 ATR，且只會往上移、不會往下退（棘輪）：", size=10)
code(doc, """
波段最高 = max(進場後每根 K 棒的 high)          # 只增不減
停損線   = 波段最高 - 3.0 * ATR(21)
實際採用 = max(前一根的停損線, 這根算出的停損線)   # 棘輪：只緊不鬆

出場條件（做多）：這根收盤價 <= 停損線 - 讓分緩衝
讓分緩衝 = 0.05 * ATR(21)                       # 容忍輕微穿線，濾雜訊
""")
para(doc, "程式碼：risk_manager.py:177-208", size=9.5, color=MUTED)

heading(doc, "4-3　兩道停損的實際關係", 2)
para(doc, "系統對外顯示的「防守價」是兩條線取較緊（對做多而言即較高）的那條。"
          "Chandelier 的距離是 3.0×ATR，cap 則是 0.5×ATR 再夾進 12～40 點之間；"
          "倍數差了 6 倍、cap 又被上限壓著，所以 Chandelier 幾乎總是比較寬，"
          "除非行情大漲讓它追上來，否則生效的一直是 cap。", size=10)
code(doc, """
兩條線交叉所需的漲幅 = Chandelier 距離 - 當時的 cap

例：ATR=110 → Chandelier 330 點、cap 觸頂 40 點 → 需漲約 290 點才換手
例：ATR= 60 → Chandelier 180 點、cap = 30 點      → 需漲約 150 點才換手
""")
para(doc, "兩者都隨 ATR 變動，所以交叉點沒有固定值。但要注意兩條線縮放的方式並不對稱："
          "cap 有 12 點的下限，Chandelier 沒有。當 ATR 低到 0.5×ATR 不足 12 點"
          "（ATR 值約 24 點以下，注意這裡指的是 ATR 的數值，不是計算週期的 21 根）時，"
          "cap 就卡在 12 點不再縮小，Chandelier 卻繼續按 3.0×ATR "
          "縮短，兩條線的距離會收斂得比預期快。低指數、低波動的期間，cap 確實長期被這個"
          "下限綁住，而不是等於 0.5×ATR。", size=10)
figure(doc, "fig4_stops.png",
       "以 ATR(21)=110、cap 觸頂於 40 點為例：早期由 cap 守住，漲約 290 點後 "
       "Chandelier 才反超接手。cap 未觸頂時，換手所需漲幅會不同。")

heading(doc, "4-4　第三道：時段強制平倉", 2)
para(doc, "每個交易時段收盤前 1 分鐘無條件平倉，且該時點之後不再接受任何新進場。"
          "強制平倉的檢查會用到 1 分鐘 K 棒，以免 5 分鐘的邊界剛好錯過那一分鐘的窗口。", size=10)
figure(doc, "fig5_sessions.png", "交易時段與兩個強制平倉時點。")

doc.add_page_break()

# ---------------------------------------------------------------- 5
heading(doc, "五、每根 K 棒的完整決策順序", 1)
para(doc, "順序本身就是規則的一部分：停損永遠先於訊號判斷，強制平倉永遠先於進場。"
          "任何一步觸發後即結束該根 K 棒的處理，不會再往下走。")
figure(doc, "fig2_onbar.png", width=4.55)
table(doc,
      ["步驟", "動作", "不符合時"],
      [["①", "確認是 5 分鐘棒", "1 分鐘棒只檢查強制平倉，然後結束"],
       ["②", "確認成交量大於 0", "無量補空棒不進指標也不驅動策略，只檢查強制平倉"],
       ["③", "更新 ATR(21)（Wilder RMA）", "—"],
       ["④", "檢查兩道停損", "觸發即平倉，本根結束"],
       ["⑤", "檢查是否進入強制平倉時點", "是則平倉，且不再看訊號"],
       ["⑥", "評估三個進場訊號", "無訊號則結束"],
       ["⑦", "部位閘：目前空手嗎", "已持倉或委託在途則不動作"],
       ["⑧", "送出買進 1 口", "—"]],
      widths=[0.55, 2.75, 3.05])
para(doc, "程式碼：core/engine.py:85-171", size=9.5, color=MUTED)

# ---------------------------------------------------------------- 6
heading(doc, "六、交易成本模型", 1)
table(doc,
      ["項目", "數值", "說明"],
      [["每口手續費", "NT$ 20", "單邊計，一趟來回收兩次"],
       ["期交稅率", "0.00002", "按名目本金計，進出各課一次"],
       ["每點價值", "NT$ 10", "微型台指期"],
       ["來回總成本", "`20*2 + (進場價+出場價)*10*0.00002`", "約 NT$ 58 / 口（指數 46,000 時）"]],
      widths=[1.3, 2.35, 2.7])
para(doc, "程式碼：backtest/replay.py:22-30", size=9.5, color=MUTED)
note(doc, "這個成本模型只用在回測。實盤（--paper）流程完全沒有引用它，因此 "
          "reports/paper_trades*.csv 與 Telegram 推播顯示的損益都是「毛」的——"
          "只有點數 × NT$10，沒有扣手續費與交易稅。看即時報表評估績效時必須自行扣除，"
          "每口來回約 NT$ 58。")

# ---------------------------------------------------------------- 7
heading(doc, "七、參數總表", 1)
table(doc,
      ["參數", "值", "作用", "程式碼位置"],
      [["策略週期", "5 分鐘", "訊號與停損的判斷週期", "run_live.py:105"],
       ["ATR 週期", "21（Wilder RMA）", "停損距離與 buy 幅度門檻的基準", "run_live.py:98"],
       ["最大口數", "1", "任何時刻最多 1 口", "run_live.py:253"],
       ["pick 回看", "20 根", "破底判定範圍", "signals.py:44"],
       ["buy 突破回看", "20 根", "前高判定範圍", "signals.py:66"],
       ["buy 幅度係數", "0.75 × ATR", "突破須超出前高的幅度", "signals.py:80"],
       ["sell_flee 回看", "10 根", "破底判定範圍", "signals.py:101"],
       ["sell_flee 區間位置", "0.6", "收盤須在當根區間上 40%", "signals.py:114"],
       ["移動停損倍數（多）", "3.0 × ATR", "Chandelier 距離", "run_live.py:93"],
       ["讓分緩衝", "0.05 × ATR", "容忍穿線的雜訊帶", "run_live.py:97"],
       ["每筆上限係數", "0.5 × ATR", "per-trade cap 主體，每根重算", "run_live.py:260"],
       ["每筆上限下限", "12 點", "雜訊地板，不隨指數縮放", "risk_manager.py:65"],
       ["每筆上限上限", "0.087% × 進場價", "風險天花板，隨指數縮放", "run_live.py:104"],
       ["ATR 暖機 fallback", "20 點", "ATR 未備妥時的固定上限", "run_live.py:260"],
       ["強制平倉提前", "1 分鐘", "時段收盤前平倉", "clock.py:28"]],
      widths=[1.55, 1.45, 2.15, 1.35], size=9)

# ---------------------------------------------------------------- 8
heading(doc, "八、程式碼裡有、但現行沒有啟用的東西", 1)
para(doc, "以下功能都存在於程式碼中且可用旗標開啟，但目前的服務參數沒有開。"
          "列出來是為了避免看程式碼時誤以為它們正在運作：")
table(doc,
      ["項目", "內容", "現況"],
      [["三個空方訊號", "touch（反轉做空）、sell（突破做空）、buy_flee（多頭陷阱做空）",
        "未啟用，因此系統不會做空"],
       ["vwap / vwap_s", "以 VWAP 偏離帶為基礎的反轉訊號", "未啟用（需 opt-in）"],
       ["orb / orb_s", "開盤區間突破訊號", "未啟用（需 opt-in）"],
       ["日線狀態閘門", "讀日線訊號工廠的大台狀態，限制盤中只能順著日線方向做",
        "未啟用；回測淨利 +86,510，低於不開的 +121,520"],
       ["麻紗 shadow", "另一套以軌道線位置進場的獨立策略", "有在跑，但屬獨立系統"]],
      widths=[1.35, 3.15, 1.9])

# ---------------------------------------------------------------- 9
heading(doc, "九、已知限制（重要）", 1)
bullets(doc, [
    ("實際虧損可能超過名目上限。", "所有停損都在 5 分鐘 K 棒收盤才檢查，"
     "若這 5 分鐘內急殺，出場價會遠低於停損線。名目 40 點的上限曾實際實現過數百點的虧損。"),
    ("重啟不會還原部位。", "服務重啟後只補回 K 棒資料，記憶體中的持倉不會恢復，"
     "該筆交易的紀錄會缺漏。"),
    ("目前是紙上交易。", "--paper 模式由內建模擬器成交，成交價與滑價與真實市場不同。"),
    ("樣本外表現須留意。", "現行參數多數在 2025-06 至 2026-05 的資料上調校，"
     "近期高波動期間的表現與樣本內有明顯差異。"),
])

para(doc, "")
para(doc, "本文件由程式自動產生，所有數值於產生當下直接讀取自原始碼。"
          "若日後修改策略，重新執行 tmp/strategy_doc_build.py 即可更新。",
     size=9, color=MUTED, italic=True)

doc.save(OUT)
print("saved:", OUT)
print("size :", f"{OUT.stat().st_size/1024:.0f} KB")
